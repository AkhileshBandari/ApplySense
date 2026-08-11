import os
import io
import tempfile
from docx import Document
from resumes.models import ResumeVersion

class ResumeRenderingService:
    @staticmethod
    def render_docx(version: ResumeVersion) -> io.BytesIO:
        """
        Renders the structured content of an APPROVED ResumeVersion into a DOCX byte stream.
        """
        if version.status != 'APPROVED':
            raise ValueError("Only APPROVED resume versions can be exported.")

        document = Document()
        content = version.structured_content

        # 1. Contact Info
        contact = content.get('contact', {})
        if contact:
            name = contact.get('name', 'Name Not Provided')
            document.add_heading(name, 0)
            
            contact_str = []
            if contact.get('email'): contact_str.append(contact['email'])
            if contact.get('phone'): contact_str.append(contact['phone'])
            if contact.get('location'): contact_str.append(contact['location'])
            if contact.get('linkedin'): contact_str.append(contact['linkedin'])
            
            if contact_str:
                p = document.add_paragraph(' | '.join(contact_str))
                p.alignment = 1 # Center aligned

        # 2. Summary
        summary = content.get('summary', '')
        if summary:
            document.add_heading('Summary', level=1)
            document.add_paragraph(summary)

        # 3. Experience
        experience = content.get('experience', [])
        if experience:
            document.add_heading('Experience', level=1)
            for exp in experience:
                title = exp.get('title', 'Role')
                company = exp.get('company', 'Company')
                dates = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                
                p = document.add_paragraph()
                p.add_run(f"{title} at {company}").bold = True
                p.add_run(f"\n{dates}")
                
                # Assume description might be a bulleted list separated by newlines
                desc = exp.get('description', '')
                if desc:
                    for bullet in desc.split('\n'):
                        if bullet.strip():
                            # Remove existing bullet points if any
                            clean_bullet = bullet.strip().lstrip('•*- ')
                            document.add_paragraph(clean_bullet, style='List Bullet')

        # 4. Education
        education = content.get('education', [])
        if education:
            document.add_heading('Education', level=1)
            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
                
                p = document.add_paragraph()
                p.add_run(f"{degree} - {institution}").bold = True
                p.add_run(f"\n{dates}")

        # 5. Skills
        skills = content.get('skills', [])
        if skills:
            document.add_heading('Skills', level=1)
            # Assuming skills is a list of strings
            document.add_paragraph(', '.join(skills))
            
        # 6. Projects
        projects = content.get('projects', [])
        if projects:
            document.add_heading('Projects', level=1)
            for proj in projects:
                name = proj.get('name', 'Project Name')
                desc = proj.get('description', '')
                p = document.add_paragraph()
                p.add_run(name).bold = True
                if desc:
                    document.add_paragraph(desc, style='List Bullet')

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
